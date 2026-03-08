import os
import re
import textwrap
from io import BytesIO
from datetime import datetime
from zoneinfo import ZoneInfo
from html import unescape

import docx
import json
import logging
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
    
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas as rl_canvas
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

logger = logging.getLogger(__name__)


def _clean_html(text: str) -> str:
    """Remove HTML tags and unescape HTML entities."""
    if not text:
        return ""
    # Remove HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    # Unescape HTML entities (&nbsp;, &lt;, etc.)
    text = unescape(text)
    return text.strip()


def _format_latex_readable(text: str) -> str:
    """
    Convert LaTeX expressions to more readable plain text format.
    Handles common patterns like fractions, superscripts, subscripts, etc.
    """
    if not text:
        return ""
    
    # Keep LaTeX delimiters visible for clarity
    text = text.replace('\\[', '\n')
    text = text.replace('\\]', '\n')
    text = text.replace('$$', '\n')
    
    # Convert common LaTeX commands to Unicode approximations
    latex_replacements = {
        r'\\frac\{([^}]+)\}\{([^}]+)\}': r'(\1)/(\2)',  # Fractions
        r'\^(\d)': lambda m: '⁰¹²³⁴⁵⁶⁷⁸⁹'[int(m.group(1))],  # Superscripts
        r'_(\d)': lambda m: '₀₁₂₃₄₅₆₇₈₉'[int(m.group(1))],  # Subscripts
        r'\\sqrt\{([^}]+)\}': r'√(\1)',  # Square root
        r'\\pi': 'π',
        r'\\alpha': 'α',
        r'\\beta': 'β',
        r'\\gamma': 'γ',
        r'\\delta': 'δ',
        r'\\theta': 'θ',
        r'\\lambda': 'λ',
        r'\\mu': 'μ',
        r'\\sigma': 'σ',
        r'\\omega': 'ω',
        r'\\Delta': 'Δ',
        r'\\Sigma': 'Σ',
        r'\\infty': '∞',
        r'\\pm': '±',
        r'\\times': '×',
        r'\\div': '÷',
        r'\\leq': '≤',
        r'\\geq': '≥',
        r'\\neq': '≠',
        r'\\approx': '≈',
        r'\\equiv': '≡',
        r'\\cdot': '·',
        r'\\rightarrow': '→',
        r'\\Rightarrow': '⇒',
        r'\\leftarrow': '←',
        r'\\Leftarrow': '⇐',
        r'\\leftrightarrow': '↔',
        r'\\partial': '∂',
        r'\\nabla': '∇',
        r'\\int': '∫',
        r'\\sum': '∑',
        r'\\prod': '∏',
        r'\\in': '∈',
        r'\\notin': '∉',
        r'\\subset': '⊂',
        r'\\subseteq': '⊆',
        r'\\cup': '∪',
        r'\\cap': '∩',
        r'\\emptyset': '∅',
        r'\\forall': '∀',
        r'\\exists': '∃',
        r'\\neg': '¬',
        r'\\wedge': '∧',
        r'\\vee': '∨',
    }
    
    # Apply simple string replacements first
    for pattern, replacement in latex_replacements.items():
        if isinstance(replacement, str):
            text = text.replace(pattern, replacement)
        else:
            text = re.sub(pattern, replacement, text)
    
    # Clean up remaining LaTeX braces
    text = text.replace('{', '').replace('}', '')
    
    # Clean up extra backslashes
    text = re.sub(r'\\([a-zA-Z]+)', r'\1', text)
    
    return text.strip()


def _safe_filename(name: str, max_len: int = 180) -> str:
    """Make a string safe for use as a filename."""
    if not name:
        return "Quiz_Results"
    s = str(name).strip()
    s = re.sub(r'\s+', '_', s)
    s = re.sub(r'[\\/:"*?<>|]+', '_', s)
    return s[:max_len] or "Quiz_Results"


def _format_quiz_text(results: dict) -> str:
    """Format quiz questions, answers, and explanations for sharing/revision (no scores)."""
    details = results.get('details', [])
    subject = results.get('subject', 'Quiz')
    difficulty = results.get('difficulty', 'Medium')
    source_filename = results.get('source_filename', '')
    
    tz = ZoneInfo('Africa/Accra')
    timestamp = datetime.now(tz).strftime('%Y-%m-%d %H:%M:%S %Z')
    
    lines = [
        f"{subject.upper()} - QUIZ QUESTIONS & ANSWERS",
        '=' * 70,
        ' '
    ]
    
    if source_filename:
        lines.append(f"Source Material: {source_filename}")
    
    lines.extend([
        f"Difficulty Level: {difficulty.title()}",
        f"Generated: {timestamp}",
        f"Total Questions: {len(details)}",
        '=' * 70,
        ' '
    ])
    
    # Add questions, answers, and explanations only (no scores or user answers)
    for idx, detail in enumerate(details, 1):
        question_text = _clean_html(detail.get('question', ''))
        question_text = _format_latex_readable(question_text)
        
        options = detail.get('options', [])
       
        # Clean and format answer
        answer_text = _clean_html(detail.get('correct_answer', 'N/A'))
        answer_text = _format_latex_readable(answer_text)
        
        # Clean and format explanation
        explanation_text = detail.get('explanation', '')
        if explanation_text:
            explanation_text = _clean_html(explanation_text)
            explanation_text = _format_latex_readable(explanation_text)
        
        # Format question block - clean and copy-paste friendly
        lines.append(f"Question {idx}: {question_text}")
       
           # Add options if this is an MCQ question
        if options:
            option_labels = ['A', 'B', 'C', 'D']
            for i, option in enumerate(options[:4]):  # Ensure max 4 options
                clean_option = _clean_html(str(option))
                clean_option = _format_latex_readable(clean_option)
                lines.append(f"   {option_labels[i]}. {clean_option}")
            lines.append('')  # Blank line after options
       
        lines.append(f"Answer: {answer_text}")
        
        if explanation_text:
            lines.append(f"Explanation: {explanation_text}")
        
        # Add spacing between questions
        lines.append('-' * 100)
        lines.append(' ')
    
    # Footer
    lines.extend([
        '',
        'This quiz was generated by Lamla AI',
        'https://lamla-ai.vercel.app',
    ])
    
    return '\n'.join(lines)


@csrf_exempt
@require_http_methods(["POST"])
def download_quiz_results(request):
    """Download quiz results as PDF, DOCX, or TXT."""
    try:
        data = json.loads(request.body) if request.body else {}
        results = data.get('results', {})
        file_format = data.get('format', 'pdf').lower()
        
        if not results:
            return JsonResponse({"error": "No results data provided"}, status=400)
        
        subject = results.get('subject', 'Quiz')
        safe_filename = _safe_filename(subject)
        tz = ZoneInfo('Africa/Accra')
        timestamp = datetime.now(tz).strftime('%Y%m%d_%H%M%S')
        filename_base = f"{safe_filename}_Quiz_{timestamp}"
        
        text_content = _format_quiz_text(results)
        
        if file_format == 'pdf':
            if not HAS_REPORTLAB:
                return JsonResponse(
                    {"error": "PDF generation not available. Please use TXT or DOCX format."},
                    status=400
                )
            
            buffer = BytesIO()
            p = rl_canvas.Canvas(buffer, pagesize=letter)
            width, height = letter
            y = height - 40
            
            for line in text_content.split('\n'):
                for wrapped_line in textwrap.wrap(line, width=95):
                    p.drawString(40, y, wrapped_line)
                    y -= 16
                    if y < 40:
                        p.showPage()
                        y = height - 40
            
            p.save()
            buffer.seek(0)
            response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{filename_base}.pdf"'
            return response
        
        elif file_format == 'docx':
            buffer = BytesIO()
            doc = docx.Document()
            
            # Add header
            header = doc.add_heading(f"{results.get('subject', 'Quiz').upper()} - QUIZ QUESTIONS & ANSWERS", level=1)
            header.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Add metadata
            if results.get('source_filename'):
                doc.add_paragraph(f"Source Material: {results.get('source_filename')}")
            doc.add_paragraph(f"Difficulty Level: {results.get('difficulty', 'Medium').title()}")
            doc.add_paragraph(f"Generated: {datetime.now(ZoneInfo('Africa/Accra')).strftime('%Y-%m-%d %H:%M:%S %Z')}")
            doc.add_paragraph(f"Total Questions: {len(results.get('details', []))}")
            doc.add_paragraph()  # Blank line
            
            # Add questions
            for idx, detail in enumerate(results.get('details', []), 1):
                q_para = doc.add_paragraph()
                q_text = _clean_html(detail.get('question', ''))
                q_text = _format_latex_readable(q_text)
                q_para.add_run(f"Question {idx}: ").bold = True
                q_para.add_run(q_text)
                
                # Add options if this is an MCQ question
                options = detail.get('options', [])
                
                if options:
                    option_labels = ['A', 'B', 'C', 'D']
                    for i, option in enumerate(options[:4]):  # Ensure max 4 options
                        opt_para = doc.add_paragraph(style='List Number')
                        clean_option = _clean_html(str(option))
                        clean_option = _format_latex_readable(clean_option)
                        opt_para.add_run(f"   {option_labels[i]}. {clean_option}")
                    doc.add_paragraph() 
               
                # Answer
                a_para = doc.add_paragraph()
                a_text = _clean_html(detail.get('correct_answer', 'N/A'))
                a_text = _format_latex_readable(a_text)
                a_para.add_run("Answer: ").bold = True
                a_para.add_run(a_text)
                
                # Explanation (if present)
                if detail.get('explanation'):
                    e_para = doc.add_paragraph()
                    e_text = _clean_html(detail.get('explanation', ''))
                    e_text = _format_latex_readable(e_text)
                    e_para.add_run("Explanation: ").bold = True
                    e_para.add_run(e_text)
                
                # Spacing between questions
                doc.add_paragraph()
            
            # Footer
            doc.add_paragraph()
            footer_para = doc.add_paragraph("This quiz was generated by Lamla AI")
            footer_para.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
            link_para = doc.add_paragraph("https://lamla-ai.vercel.app")
            link_para.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
            
            doc.save(buffer)
            buffer.seek(0)
            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{filename_base}.docx"'
            return response
        
        else:  # TXT
            response = HttpResponse(text_content, content_type='text/plain')
            response['Content-Disposition'] = f'attachment; filename="{filename_base}.txt"'
            return response
    
    except Exception as e:
        logger.error(f"Error generating download: {e}", exc_info=True)
        return JsonResponse({"error": "Failed to generate file"}, status=500)

